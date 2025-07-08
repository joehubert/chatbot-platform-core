"""
Document Processing Service

This module handles document text extraction, chunking, and embedding generation
for various file formats including PDF, DOCX, TXT, and Markdown.
"""

import asyncio
import logging
from typing import List, Dict, Any, Optional, Union
from dataclasses import dataclass
from datetime import datetime
import hashlib
import tempfile
import os
from pathlib import Path

# Text extraction libraries
try:
    import PyPDF2
    import pdfplumber
except ImportError:
    PyPDF2 = None
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import markdown
except ImportError:
    markdown = None

# Embedding libraries
try:
    import openai
except ImportError:
    openai = None

try:
    from sentence_transformers import SentenceTransformer
except ImportError:
    SentenceTransformer = None

logger = logging.getLogger(__name__)


@dataclass
class DocumentChunk:
    """A chunk of processed document content"""
    content: str
    embedding: List[float]
    metadata: Dict[str, Any]
    start_char: int
    end_char: int
    chunk_index: int


@dataclass
class ProcessedDocument:
    """Processed document with chunks and metadata"""
    content: str
    chunks: List[DocumentChunk]
    metadata: Dict[str, Any]
    processing_time_ms: float


class DocumentProcessor:
    """Document processing service with text extraction and chunking"""
    
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        embedding_model: str = "sentence-transformers/all-MiniLM-L6-v2",
        openai_api_key: Optional[str] = None
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.embedding_model_name = embedding_model
        self.openai_api_key = openai_api_key
        self.embedding_model = None
        self._initialized = False
        
        # Supported file types
        self.supported_types = {
            'application/pdf': self._extract_pdf_text,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_docx_text,
            'text/plain': self._extract_txt_text,
            'text/markdown': self._extract_markdown_text,
            'application/msword': self._extract_docx_text  # Legacy DOC files
        }
    
    async def initialize(self) -> bool:
        """Initialize the document processor"""
        try:
            # Initialize embedding model
            if "openai" in self.embedding_model_name.lower() and self.openai_api_key:
                # Use OpenAI embeddings
                if openai is None:
                    logger.error("OpenAI package not installed. Install with: pip install openai")
                    return False
                openai.api_key = self.openai_api_key
                self.embedding_model = "openai"
                logger.info("Initialized with OpenAI embeddings")
            else:
                # Use Sentence Transformers
                if SentenceTransformer is None:
                    logger.error("SentenceTransformers package not installed. Install with: pip install sentence-transformers")
                    return False
                
                # Load model in thread to avoid blocking
                self.embedding_model = await self._load_sentence_transformer()
                if self.embedding_model is None:
                    return False
                logger.info(f"Initialized with SentenceTransformers model: {self.embedding_model_name}")
            
            self._initialized = True
            return True
            
        except Exception as e:
            logger.error(f"Failed to initialize document processor: {str(e)}")
            return False
    
    async def _load_sentence_transformer(self):
        """Load SentenceTransformer model in a thread"""
        try:
            loop = asyncio.get_event_loop()
            model = await loop.run_in_executor(
                None, 
                SentenceTransformer, 
                self.embedding_model_name
            )
            return model
        except Exception as e:
            logger.error(f"Failed to load SentenceTransformer model: {str(e)}")
            return None
    
    def _ensure_initialized(self):
        """Ensure the processor is initialized"""
        if not self._initialized:
            raise RuntimeError("Document processor not initialized. Call initialize() first.")
    
    async def process_document(
        self,
        file_content: bytes,
        filename: str,
        content_type: str
    ) -> Optional[ProcessedDocument]:
        """Process a document: extract text, chunk, and generate embeddings"""
        self._ensure_initialized()
        
        start_time = datetime.utcnow()
        
        try:
            # Extract text from document
            text_content = await self._extract_text(file_content, content_type)
            
            if not text_content or not text_content.strip():
                logger.warning(f"No text content extracted from {filename}")
                return None
            
            # Create chunks
            chunks = await self._create_chunks(text_content)
            
            if not chunks:
                logger.warning(f"No chunks created from {filename}")
                return None
            
            # Generate embeddings for each chunk
            processed_chunks = []
            for i, chunk_text in enumerate(chunks):
                try:
                    embedding = await self.generate_embedding(chunk_text)
                    
                    if embedding:
                        # Calculate character positions
                        start_char = sum(len(c) for c in chunks[:i])
                        end_char = start_char + len(chunk_text)
                        
                        chunk = DocumentChunk(
                            content=chunk_text,
                            embedding=embedding,
                            metadata={
                                'filename': filename,
                                'content_type': content_type,
                                'chunk_method': 'fixed_size_overlap'
                            },
                            start_char=start_char,
                            end_char=end_char,
                            chunk_index=i
                        )
                        processed_chunks.append(chunk)
                    else:
                        logger.warning(f"Failed to generate embedding for chunk {i} in {filename}")
                        
                except Exception as e:
                    logger.error(f"Error processing chunk {i} in {filename}: {str(e)}")
            
            if not processed_chunks:
                logger.error(f"No chunks with embeddings created for {filename}")
                return None
            
            # Calculate processing time
            processing_time = (datetime.utcnow() - start_time).total_seconds() * 1000
            
            # Create processed document
            doc_metadata = {
                'filename': filename,
                'content_type': content_type,
                'original_length': len(text_content),
                'chunk_count': len(processed_chunks),
                'chunk_size': self.chunk_size,
                'chunk_overlap': self.chunk_overlap,
                'embedding_model': self.embedding_model_name,
                'processed_at': datetime.utcnow().isoformat()
            }
            
            return ProcessedDocument(
                content=text_content,
                chunks=processed_chunks,
                metadata=doc_metadata,
                processing_time_ms=processing_time
            )
            
        except Exception as e:
            logger.error(f"Failed to process document {filename}: {str(e)}")
            return None
    
    async def _extract_text(self, file_content: bytes, content_type: str) -> str:
        """Extract text from file content based on content type"""
        extractor = self.supported_types.get(content_type)
        
        if not extractor:
            raise ValueError(f"Unsupported content type: {content_type}")
        
        return await extractor(file_content)
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF content"""
        if PyPDF2 is None and pdfplumber is None:
            raise RuntimeError("PDF libraries not installed. Install with: pip install PyPDF2 pdfplumber")
        
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name
            
            try:
                # Try pdfplumber first (better text extraction)
                if pdfplumber:
                    text = await self._extract_with_pdfplumber(tmp_file_path)
                    if text.strip():
                        return text
                
                # Fallback to PyPDF2
                if PyPDF2:
                    text = await self._extract_with_pypdf2(tmp_file_path)
                    return text
                
                return ""
                
            finally:
                # Clean up temporary file
                os.unlink(tmp_file_path)
                
        except Exception as e:
            logger.error(f"Failed to extract PDF text: {str(e)}")
            return ""
    
    async def _extract_with_pdfplumber(self, file_path: str) -> str:
        """Extract text using pdfplumber"""
        def extract():
            text_parts = []
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n\n".join(text_parts)
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, extract)
    
    async def _extract_with_pypdf2(self, file_path: str) -> str:
        """Extract text using PyPDF2"""
        def extract():
            text_parts = []
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            return "\n\n".join(text_parts)
        
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, extract)
    
    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX content"""
        if DocxDocument is None:
            raise RuntimeError("python-docx not installed. Install with: pip install python-docx")
        
        try:
            # Create temporary file
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name
            
            try:
                def extract():
                    doc = DocxDocument(tmp_file_path)
                    text_parts = []
                    
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_parts.append(paragraph.text)
                    
                    return "\n\n".join(text_parts)
                
                loop = asyncio.get_event_loop()
                return await loop.run_in_executor(None, extract)
                
            finally:
                # Clean up temporary file
                os.unlink(tmp_file_path)
                
        except Exception as e:
            logger.error(f"Failed to extract DOCX text: {str(e)}")
            return ""
    
    async def _extract_txt_text(self, file_content: bytes) -> str:
        """Extract text from plain text content"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            
            for encoding in encodings:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, use utf-8 with error replacement
            return file_content.decode('utf-8', errors='replace')
            
        except Exception as e:
            logger.error(f"Failed to extract text content: {str(e)}")
            return ""
    
    async def _extract_markdown_text(self, file_content: bytes) -> str:
        """Extract text from Markdown content"""
        try:
            # First get the text content
            text_content = await self._extract_txt_text(file_content)
            
            if markdown:
                # Convert markdown to HTML then extract text
                # This preserves structure better than raw markdown
                def convert_markdown():
                    html = markdown.markdown(text_content)
                    # Simple HTML tag removal (for better text extraction)
                    import re
                    text = re.sub(r'<[^>]+>', '', html)
                    return text
                
                loop = asyncio.get_event_loop()
                return await loop.run_in_executor(None, convert_markdown)
            else:
                # If markdown library not available, return raw text
                return text_content
                
        except Exception as e:
            logger.error(f"Failed to extract Markdown text: {str(e)}")
            return ""
    
    async def _create_chunks(self, text: str) -> List[str]:
        """Create text chunks with overlap"""
        try:
            if len(text) <= self.chunk_size:
                return [text]
            
            chunks = []
            start = 0
            
            while start < len(text):
                # Calculate end position
                end = start + self.chunk_size
                
                # If this is not the last chunk, try to break at word boundary
                if end < len(text):
                    # Look for word boundary within overlap distance
                    boundary_search_start = max(end - self.chunk_overlap, start + 1)
                    
                    # Find the last space or newline in the search area
                    for i in range(end, boundary_search_start - 1, -1):
                        if text[i] in [' ', '\n', '.', '!', '?']:
                            end = i + 1
                            break
                
                # Extract chunk
                chunk = text[start:end].strip()
                if chunk:
                    chunks.append(chunk)
                
                # Move start position for next chunk
                if end >= len(text):
                    break
                
                start = end - self.chunk_overlap
                
                # Ensure we don't go backwards
                if start <= len(chunks[-1]) + (len(chunks) - 1) * (self.chunk_size - self.chunk_overlap):
                    start = end
            
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to create chunks: {str(e)}")
            return []
    
    async def generate_embedding(self, text: str) -> Optional[List[float]]:
        """Generate embedding for text"""
        self._ensure_initialized()
        
        try:
            if self.embedding_model == "openai":
                return await self._generate_openai_embedding(text)
            else:
                return await self._generate_sentence_transformer_embedding(text)
                
        except Exception as e:
            logger.error(f"Failed to generate embedding: {str(e)}")
            return None
    
    async def _generate_openai_embedding(self, text: str) -> Optional[List[float]]:
        """Generate embedding using OpenAI"""
        try:
            def get_embedding():
                response = openai.Embedding.create(
                    model="text-embedding-ada-002",
                    input=text
                )
                return response['data'][0]['embedding']
            
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, get_embedding)
            
        except Exception as e:
            logger.error(f"OpenAI embedding failed: {str(e)}")
            return None
    
    async def _generate_sentence_transformer_embedding(self, text: str) -> Optional[List[float]]:
        """Generate embedding using SentenceTransformers"""
        try:
            def get_embedding():
                embedding = self.embedding_model.encode([text])
                return embedding[0].tolist()
            
            loop = asyncio.get_event_loop()
            return await loop.run_in_executor(None, get_embedding)
            
        except Exception as e:
            logger.error(f"SentenceTransformer embedding failed: {str(e)}")
            return None
    
    def get_supported_types(self) -> List[str]:
        """Get list of supported content types"""
        return list(self.supported_types.keys())
    
    def is_supported_type(self, content_type: str) -> bool:
        """Check if content type is supported"""
        return content_type in self.supported_types
    
    async def get_text_preview(self, file_content: bytes, content_type: str, max_chars: int = 500) -> str:
        """Get a preview of the text content without full processing"""
        try:
            text = await self._extract_text(file_content, content_type)
            if len(text) <= max_chars:
                return text
            return text[:max_chars] + "..."
        except Exception as e:
            logger.error(f"Failed to get text preview: {str(e)}")
            return ""
    
    async def health_check(self) -> Dict[str, Any]:
        """Perform health check on the document processor"""
        try:
            if not self._initialized:
                return {"status": "unhealthy", "reason": "not_initialized"}
            
            # Test embedding generation
            test_embedding = await self.generate_embedding("test text")
            
            if test_embedding:
                return {
                    "status": "healthy",
                    "embedding_model": self.embedding_model_name,
                    "supported_types": len(self.supported_types),
                    "embedding_dimension": len(test_embedding)
                }
            else:
                return {"status": "unhealthy", "reason": "embedding_generation_failed"}
                
        except Exception as e:
            return {
                "status": "unhealthy",
                "reason": "exception",
                "error": str(e)
            }


# Factory function for easy service creation
def create_document_processor(
    chunk_size: int = 1000,
    chunk_overlap: int = 200,
    embedding_model: str = "sentence-transformers/all-MiniLM-L6-v2",
    openai_api_key: Optional[str] = None
) -> DocumentProcessor:
    """Factory function to create a document processor"""
    return DocumentProcessor(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        embedding_model=embedding_model,
        openai_api_key=openai_api_key
    )